import streamlit as st
from datetime import datetime

st.set_page_config(page_title="Motorin Screener", layout="centered")

# ---- LOGO (CENTERED AND ENLARGED) ----
st.markdown("<div style='text-align: center;'><img src='https://i.imgur.com/1thkHWE.png' width='480'></div>", unsafe_allow_html=True)
st.markdown("---")

# ---- DEMOGRAPHICS ----
st.header("Screening Details")
child_first_name = st.text_input("Child's First Name")
child_last_name = st.text_input("Child's Last Name")
dob = st.date_input("Date of Birth")
date_of_screen = st.date_input("Date of Screen", datetime.today())
therapist_name = st.text_input("Therapist Name")

# ---- CHRONOLOGICAL AGE CALCULATION ----
if dob:
    today = date_of_screen
    delta = today - dob
    years = delta.days // 365
    months = (delta.days % 365) // 30
    st.markdown(f"**Chronological Age:** {years} years, {months} months")

st.markdown("---")

# ---- SCREENER INSTRUCTIONS ----
st.subheader("Instructions")
st.markdown("Rate each item based on observation or parent/therapist report using the following scale:")
st.markdown("- 0 = Absent")
st.markdown("- 1 = Emerging")
st.markdown("- 2 = Present")
st.markdown("---")

# ---- MOTORIN SCREENER ITEMS ----
motorin_items = {
    "0-3 Months": [
        "Tracks rattle side to side",
        "Brings hands to midline",
        "Briefly brings hands to mouth",
        "Grips caregiver's finger",
        "Opens fingers during touch/play",
        "Shows alertness to visual stimuli"
    ],
    "4-6 Months": [
        "Reaches for toy with both hands",
        "Holds a rattle briefly with one or both hands",
        "Brings toy to mouth",
        "Swipes at toy dangling in front of them",
        "Maintains grasp when toy is gently pulled",
        "Purposefully turns head to look at sounds"
    ],
    "7-9 Months": [
        "Transfers toy between hands",
        "Bangs two toys together",
        "Rakes small object with fingers",
        "Grasps cube with radial-palmar grasp",
        "Shakes toy purposefully",
        "Reaches in different directions with control"
    ],
    "10-12 Months": [
        "Releases block into container",
        "Uses a pincer grasp to pick up tiny objects",
        "Opens hand to drop object without cue",
        "Pokes with index finger",
        "Takes rings off of a stacker"
    ],
    "13-18 Months": [
        "Turns pages in cardboard book (multiple at once)",
        "Scribbles with crayon (full-arm motion)",
        "Places 1 shape into shape sorter",
        "Removes pegs from pegboard",
        "Places toy in/out of container",
        "Uses both hands in midline play",
        "Puts 4 rings onto a stacker"
    ],
    "19-24 Months": [
        "Builds tower of 4+ blocks",
        "Places lid on container",
        "Unscrews or pulls off container lid",
        "Scribbles spontaneously (in all directions)",
        "Places 3+ shapes in sorter",
        "Inserts coin into slot",
        "Takes apart pop beads",
        "Imitates forming vertical strokes"
    ],
    "25-36 Months (2-3 Years)": [
        "Draws vertical stroke",
        "Builds 6-block tower",
        "Opens simple flip-top container",
        "Strings 1–2 large beads",
        "Imitates circular scribble",
        "Begins to show hand preference",
        "Puts together pop beads",
        "Imitates forming horizontal strokes"
    ],
    "37-48 Months (3-4 Years)": [
        "Snips paper with child scissors",
        "Imitates forming a circle",
        "Strings 4–5 beads",
        "Cuts across paper with moderate control",
        "Builds 9–10 block tower",
        "Screws lid onto container",
        "Imitates forming a cross"
    ],
    "49-60 Months (4-5 Years)": [
        "Imitates forming a square",
        "Draws simple person (head + limbs)",
        "Cuts on thick line with control",
        "Colors inside simple shape (e.g., circle)",
        "Uses mature tripod grasp (emerging)",
        "Places clothespin onto object"
    ],
    "61-72 Months (5-6 Years)": [
        "Imitates forming a triangle",
        "Draws complete person (head, limbs, facial features)",
        "Cuts simple shapes (circle, square)",
        "Writes name or letters with legibility",
        "Strings 3 small beads onto a string",
        "Imitates and reproduces 3-step visual patterns",
        "Opens food packages"
    ],
    "73-84 Months (6-7 Years)": [
        "Copies diamond",
        "Writes full name with legible spacing",
        "Cuts out a square accurately",
        "Writes short sentence",
        "Folds paper in half with symmetry",
        "Completes dot-to-dot picture with precision"
    ]
}

responses = {}

for band, items in motorin_items.items():
    st.subheader(band)
    for item in items:
        key = f"{band}_{item}"
        response = st.radio(
            label=item,
            options=["Absent", "Emerging", "Present"],
            key=key,
            horizontal=True
        )
        responses[key] = response

if st.button("Submit"):
    st.success("Screening complete! Generating summary...")
    passed = sum(1 for r in responses.values() if r == "Present")
    total = len(responses)
    st.write(f"Items marked 'Present': {passed} / {total}")

    # Generate paragraph summary for families
    summary = f"""
    Based on today's screening using the MOTORIN tool, {child_first_name} {child_last_name} demonstrated {passed} of {total} fine motor skills as developmentally present. These observed skills include a variety of age-appropriate abilities such as grasping, drawing, manipulating objects, and imitating motor patterns. 

    This information provides a helpful snapshot of {child_first_name}'s current strengths and emerging abilities. It can be used to support goal planning, therapy recommendations, or simply to inform caregivers of developmental progress. Continued observation and practice of fine motor tasks at home can encourage skill development.
    """

    st.markdown("---")
    st.subheader("Family-Friendly Summary")
    st.markdown(summary)

    # Save to Word and PDF
    from docx import Document
    from fpdf import FPDF

    doc = Document()
    doc.add_heading("MOTORIN Family Summary", level=1)
    doc.add_paragraph(summary.strip())
    word_filename = f"motorin_summary_{child_first_name}_{child_last_name}.docx"
    doc.save(word_filename)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in summary.strip().split("
"):
        pdf.multi_cell(0, 10, line)
    pdf_filename = f"motorin_summary_{child_first_name}_{child_last_name}.pdf"
    pdf.output(pdf_filename)

    with open(word_filename, "rb") as f:
        st.download_button("Download Word Summary", f, file_name=word_filename)

    with open(pdf_filename, "rb") as f:
        st.download_button("Download PDF Summary", f, file_name=pdf_filename)
